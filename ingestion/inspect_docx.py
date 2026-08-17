from pathlib import Path
from docx import Document


DOCX_PATH = "data/raw/TS 23.501/23501-k20/23501-k20.docx"


def inspect_docx(path: Path):
    document = Document(path)

    print(f"Paragraphs: {len(document.paragraphs)}")
    print(f"Tables: {len(document.tables)}")

    print("\n--- FIRST 50 PARAGRAPHS ---\n")

    for i, paragraph in enumerate(document.paragraphs[:50]):
        text = paragraph.text.strip()

        if text:
            print(
                f"[{i}] "
                f"style={paragraph.style.name!r} "
                f"text={text[:300]!r}"
            )


if __name__ == "__main__":
    inspect_docx(DOCX_PATH)