from pathlib import Path
from typing import List, Dict

from docx import Document

from ingestion.section_parser import build_sections
from ingestion.chunker import create_chunks
from ingestion.metadata import extract_document_metadata


def load_docx(docx_path: Path) -> List[Dict]:
    """
    Load a 3GPP DOCX and convert its body into section records.
    """
    document = Document(docx_path)

    sections = build_sections(document.paragraphs)

    records = []

    for section in sections:
        content = "\n".join(section.paragraphs).strip()

        if not content:
            continue

        records.append(
            {
                "section": section.number,
                "section_title": section.title,
                "level": section.level,
                "content": content,
                "start_paragraph": section.start_paragraph,
                "end_paragraph": section.end_paragraph,
            }
        )

    return records


if __name__ == "__main__":
    path = Path(
        r"data\raw\TS 23.501\23501-k20\23501-k20.docx"
    )

    # Step 1: Parse DOCX into sections
    records = load_docx(path)

    print(f"Detected sections: {len(records)}")

    # Step 2: Extract document metadata
    metadata = extract_document_metadata(path)

    print("\nDocument metadata:")
    for key, value in metadata.items():
        print(f"{key}: {value}")

    # Step 3: Convert sections into chunks
    chunks = create_chunks(
        records,
        spec_number=metadata["specification"],
        version="unknown",
        release=metadata["release"],
    )

    # Step 4: Show chunk statistics
    print(f"\nTotal chunks: {len(chunks)}")

    # Step 5: Preview first 5 chunks
    for chunk in chunks[:5]:
        print("\n" + "=" * 80)
        print(f"Chunk ID: {chunk['chunk_id']}")
        print(
            f"Section: {chunk['section']} "
            f"{chunk['section_title']}"
        )
        print(f"Content type: {chunk['content_type']}")
        print("-" * 80)
        print(chunk["content"][:500])